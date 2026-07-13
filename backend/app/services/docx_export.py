"""Sprint 4.6a (D17) — Markdown report → .docx bytes, via python-docx.

The deliverable layer is the product (PHASE4's value-proposition framing): a
real downloadable file, not chat text to copy-paste. python-docx is a small
pure-Python dependency, imported lazily by main.py's download endpoint so it
costs nothing on every other request.

Line parsing lives in report_markdown.markdown_blocks() — shared with the PDF
exporter (fix batch #3) so both downloads render the exact same structure.
Figures (Sprint 4.6b) arrive as [[figure:N]] markers + validated specs: each
renders as an embedded PNG chart (figures.render_figure_png), or as the
chart's data in plain text when rendering is unavailable — the numbers always
reach the reader.
"""
import io

from docx import Document
from docx.shared import Inches, Pt, RGBColor

from app.services.report_markdown import BOLD_SPLIT, markdown_blocks

# Dark amber, readable on the white page — the disclaimer must be seen, not
# skimmed past (it is part of the design, per the 4.6 owner clarification).
_DISCLAIMER_COLOR = RGBColor(0x92, 0x5A, 0x0A)
_MUTED_COLOR = RGBColor(0x89, 0x87, 0x81)


def _add_runs_with_bold(paragraph, text: str):
    """Split '**bold** rest' into styled runs. Unmatched ** stays literal."""
    pos = 0
    for match in BOLD_SPLIT.finditer(text):
        if match.start() > pos:
            paragraph.add_run(text[pos:match.start()])
        paragraph.add_run(match.group(1)).bold = True
        pos = match.end()
    if pos < len(text):
        paragraph.add_run(text[pos:])


def _add_figure(doc, spec: dict):
    """One validated chart spec → embedded PNG, or its data as plain text."""
    from app.services.figures import figure_fallback_lines, render_figure_png

    png = render_figure_png(spec)
    if png:
        doc.add_picture(io.BytesIO(png), width=Inches(6.0))
        caption = doc.add_paragraph()
        run = caption.add_run(f"Figure: {spec['title']}")
        run.italic = True
        run.font.size = Pt(8.5)
        run.font.color.rgb = _MUTED_COLOR
        return
    for line in figure_fallback_lines(spec):
        paragraph = doc.add_paragraph()
        run = paragraph.add_run(line)
        run.font.size = Pt(10)


def markdown_report_to_docx(report_md: str, title: str, disclaimer: str,
                            generated_note: str, figures: list[dict] | None = None) -> bytes:
    doc = Document()
    figures = figures or []

    doc.add_heading(title, level=0)

    # Disclaimer directly under the title: italic, colored, unmissable.
    disclaimer_para = doc.add_paragraph()
    run = disclaimer_para.add_run(f"⚠ {disclaimer}")
    run.italic = True
    run.font.size = Pt(10)
    run.font.color.rgb = _DISCLAIMER_COLOR

    first_heading_skipped = False
    for kind, text in markdown_blocks(report_md):
        if kind == "h1":
            # The generator opens with a "# title" line that duplicates the
            # document heading above — drop only that first one.
            if not first_heading_skipped:
                first_heading_skipped = True
                continue
            doc.add_heading(text, level=1)
        elif kind == "h2":
            doc.add_heading(text, level=1)
        elif kind == "h3":
            doc.add_heading(text, level=2)
        elif kind == "bullet":
            _add_runs_with_bold(doc.add_paragraph(style="List Bullet"), text)
        elif kind == "number":
            _add_runs_with_bold(doc.add_paragraph(style="List Number"), text)
        elif kind == "figure":
            index = int(text) - 1
            if 0 <= index < len(figures):
                _add_figure(doc, figures[index])
        else:
            _add_runs_with_bold(doc.add_paragraph(), text)

    footer = doc.add_paragraph()
    footer_run = footer.add_run(generated_note)
    footer_run.italic = True
    footer_run.font.size = Pt(9)

    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()
